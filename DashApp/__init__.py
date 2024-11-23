import pandas as pd
import dash
import dash_table
from dash.dependencies import Input, Output, State
from dash import dcc, html
#from temp import USER_PASS_MAPPING
import random
import string
from datetime import datetime
from docx import Document
from docx.shared import Pt
from supabase import create_client, Client
import os
from flask_login.utils import login_required






def create_dash_application(flask_app):
    from DashApp.temp import dfx, df, bird_requirement_data, vitamins_minerals_data, columns


    dash_app = dash.Dash(server = flask_app, name = "Feedeyes", url_base_pathname = "/dash/")
    
    dash_app.layout = html.Div(children=[ 
    html.Link(rel='stylesheet', href='/static/assets/styles_sheet.css'),  
    html.Div(children=[
        html.Div(id="sidebar", children=[
            dcc.Dropdown(
                id='stages_dd',
                options=[
                    {'label': 'Pre-starter', 'value': 'Pre-starter'},
                    {'label': 'Starter', 'value': 'Starter'},
                    {'label': 'Finisher', 'value': 'Finisher'},
                
                ], value = None, placeholder='Select Feed Phase', className= 'select'
                
            ),

            dcc.Input(id= "nBird", type= "number", debounce= False, 
                      placeholder= "Number of birds...", 
                      className= "birdN"
                      ),

            dcc.RadioItems(
                        id='table-toggle',
                        options=[
                            {'label': 'Bird Requirement', 'value': 'bird'},
                            {'label': 'Vitamins & Micro Minerals', 'value': 'vitamins'}
                        ],
                        value='bird',  className= "requirement" 
                
                    ),


            dcc.Dropdown(
                id='ingredient_dd',
                options=[{'label': ingredient, 'value': ingredient} for ingredient in dfx['INGREDIENT'].unique()],
                multi=True, placeholder="Select ingredients...",
                optionHeight= 35,
                searchable= True,
                clearable= True,
                className='custom-dropdown'  
            ),

            html.Div(
                    children=[
                        html.A("Logout", href="/logout", className="logout-button")  
                    ],
                    className="logout-container"  
                )
                        ], className="sidebox"),

         html.Div(children=[
            html.Div(id="topbar", children= [
                html.Div(children = [
                    html.Label("Feedeyes")], className="feed-analysis", 
                    ),
                    html.Div(children = [

                            html.Div(id= "ME_CP", children=[
                                html.P(html.B("ME/CP(KG)")),
                                    html.P(html.B("0"))
                                ], className='card'),
                        

                            html.Div(id = "Total_Cost",
                                children=[
                                    html.P(html.B("TOTAL COST")),
                                    html.P(html.B("0"))
                                ], className='card'),

                            html.Div(id= "COST_25kg", children=[
                                html.P(html.B("COST/25kg")),
                                    html.P(html.B("0"))
                                ], className='card'),

                     

                        html.Div([
                                dcc.Store(id='feedNeed_store'),  
                                html.Div(
                                    id="Feed_req", 
                                    children=[
                                        html.P(html.B("Feed Req.", )),
                                        html.P(html.B("0",)),
                                        html.P("#",)
                                    ], className= 'card card-3'

                                )
                            ], )
,
                                

                            html.Div(children=[
                                html.Img(src='/static/assets/chickenFeed.png')
                            ], className='chicken__id')
                            ], className= 'KPI')
                    ], className='parent-Top'),


            html.Div(children=[
                html.Div(
                    id="firstbar", children=[
                        dash_table.DataTable(
                            id='ingredient_table',
                            columns=[
                                    {
                                    'name': 'QTY(100KG)' if col == 'QUANTITY' else 'PRICE' if col == 'QUANTITY PRICE' else col, 
                                    'id': col, 
                                    'editable': col in ['QUANTITY', 'PRICE/KG']
                                }
                                for col in columns
                            ],

                            data=[],  # Initially empty
                            style_table={
                                'overflowY': 'auto',  # Enable vertical scroll
                                'overflowX': 'auto',  # Enable horizontal scroll
                                'height': '100%',     # Ensure the table fills the container height
                                'width': '100%',      # Ensure the table fills the container width
                            },
                            style_cell={
                                'minWidth': '100px', 'width': '100px', 'maxWidth': '100px',
                                'whiteSpace': 'normal',
                                'textAlign': 'left',
                                'padding': '5px'
                            },
                            style_data={
                                'height': 'auto',  # Allow rows to adjust height but limit it below
                                'maxHeight': '65px',  # Set a max height for each row
                                'overflow': 'hidden',  # Hide overflow to prevent table expansion
                            },
                            page_current=0,  # Start on the first page
                            page_size=6,     # Six rows per page
                            page_action="native",  
                            style_header={
                                'backgroundColor': 'lightgrey',
                                'fontWeight': 'bold'
                            }
                        ),

                        html.Div(id='table-container', children=[
                        html.Div(id='bird-requirement-table', children=[
                        html.H1("Bird Requirement", style={'textAlign': 'center'}),  # Fixed here
                        dash_table.DataTable(
                        columns=[
                            {"name": "TYPE OF BIRD", "id": "TYPE OF BIRD", "type": "text"},
                            {"name": "AGE", "id": "AGE", "type": "text"},
                            {"name": "ME/CP(LB)", "id": "ME/CP(LB)", "type": "numeric"},
                            {"name": "ME/CP(KG)", "id": "ME/CP(KG)", "type": "numeric"},
                            {"name": "ME", "id": "ME", "type": "numeric"},
                            {"name": "CP", "id": "CP", "type": "numeric"},
                            {"name": "ME/CP", "id": "ME/CP", "type": "numeric"}
                        ],
                        data=bird_requirement_data,
                        style_cell={'textAlign': 'center', 'fontFamily': 'Arial', 'fontSize': 14, 
                                    'minWidth': '100px', 'width': '100px', 'maxWidth': '100px'},
                        style_header={'backgroundColor': 'black', 'fontWeight': 'bold', 
                                    'color': 'white', 'textAlign': 'center'},
                        style_data_conditional=[
                            {'if': {'column_id': 'ME/CP(LB)'}, 'color': 'blue', 'fontWeight': 'bold'},
                            {'if': {'column_id': 'TYPE OF BIRD'}, 'color': 'red', 'fontWeight': 'bold'}
                        ],
                        style_table={'height': '300px', 'overflowY': 'auto', 'overflowX': 'auto'},
                        fixed_rows={'headers': True},
                        filter_action="native",
                        sort_action="native",
                        row_deletable=True,
                        page_size=6,
                        page_current=0
                        )
                    ], style={'display': 'block'}),

                    html.Div(id='vitamins-minerals-table-container', children=[
                    html.H1("Vits. & Mins. Requirement"),
                    dash_table.DataTable(
                        id='vitamins-minerals-table',
                        columns=[
                            {"name": "Category", "id": "Category"},
                            {"name": "Name", "id": "Name"},
                            {"name": "Source", "id": "Source"},
                            {"name": "Potency (Min)", "id": "Potency (Min)"},
                        ],
                        data=vitamins_minerals_data,
                        filter_action="native",
                        sort_action="native",
                        page_size=10,
                        style_header={'backgroundColor': 'black', 'fontWeight': 'bold', 
                                    'color': 'white', 'textAlign': 'center'},
                        style_cell={'textAlign': 'center', 'fontFamily': 'Arial', 
                                    'fontSize': 14, 'minWidth': '100px',
                                    'width': '100px', 'maxWidth': '100px'},
                        style_data_conditional=[
                            {'if': {'column_id': 'Name'}, 'color': 'blue', 'fontWeight': 'bold'},
                            {'if': {'column_id': "Potency (Min)"}, 'color': 'red', 'fontWeight': 'bold'}
                        ],
                        style_data={
                            'whiteSpace': 'normal',
                            'height': 'auto',
                        },
                        style_table={'overflowX': 'auto'}
                        )
                    ], style={'display': 'none'})  # Hidden by default
                    ])
                                    
                                                ]),

                    html.Div(
                    id="secondbar",
                    style={
                        'height': '475px',
                        'overflowY': 'auto',
                        'padding': '10px',
                        'border': '1px solid #ddd',
                        'boxSizing': 'border-box',
                        },

                    
                    children=[
                        # Date and Time
                        html.Div(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", style={'fontWeight': 'bold', 'marginBottom': '10px'}),


                        # Company Name Input
                        html.Div([
                            html.Label("Company Name:", style={'fontWeight': 'bold'}),
                            dcc.Input(id='Company_name', type='text', placeholder='Enter Company name', style={'marginLeft': '10px'})
                        ], style={'marginBottom': '10px'}),

                        # Feed Name Input
                        html.Div([
                            html.Label("Feed Name:", style={'fontWeight': 'bold'}),
                            dcc.Input(id='feed_name', type='text', placeholder='Enter feed name', style={'marginLeft': '10px'})
                        ], style={'marginBottom': '10px'}),

                        # Feed Code Display
                        html.Div([
                            html.Label("Feed Code:", style={'fontWeight': 'bold'}),
                            html.Span(id='feed_code', style={'marginLeft': '10px'})
                        ], style={'marginBottom': '10px'}),

                        # Feed Analysis Report Heading
                        html.H2("Feed Analysis Report", style={'fontWeight': 'bold'}),

                        # Table for selected ingredients
                        dash_table.DataTable(
                            id='report_table',
                            columns=[{'name': col, 'id': col} for col in ['INGREDIENT', 'PRICE/KG', 'QUANTITY', 'AMOUNT']],
                            data=[],  # This will be filled dynamically
                            style_table={
                                'width': '100%',
                            },
                            style_cell={
                                'textAlign': 'left',
                                'padding': '5px'
                            },
                            style_header={
                                'backgroundColor': 'lightgrey',
                                'fontWeight': 'bold'
                            }
                        ),

                        # Table for nutrient composition
                        dash_table.DataTable(
                            id='nutrient_table',
                            columns=[{'name': 'Nutrient Composition', 'id': 'Nutrient'},
                                     {'name': 'Actual', 'id': 'Actual'}],
                            data=[],  # This will be filled dynamically
                            style_table={
                                'width': '100%',
                            },
                            style_cell={
                                'textAlign': 'left',
                                'padding': '5px'
                            },
                            style_header={
                                'backgroundColor': 'lightgrey',
                                'fontWeight': 'bold'
                            }
                        ),

                        #print button
                        html.Button("Download", id="download_word_button", n_clicks=0, style={'marginTop': '20px', 'fontWeight': 'bold'}),
                        dcc.Download(id="download_word")

                    ]
                    )
                    ], className="rectangle")
                    ], className="box_arrange")
                    ], className="all_box")
                    ], className="container")



    @dash_app.callback(
        Output('feed_code', 'children'),
        [Input('feed_name', 'value')]
    )
    def generate_feed_code(feed_name):
        return ''.join(random.choices(string.ascii_uppercase + string.digits, k=4))



    @dash_app.callback(
        [Output('bird-requirement-table', 'style'),
        Output('vitamins-minerals-table-container', 'style')],
        [Input('table-toggle', 'value')]
    )
    def toggle_tables(selected_value):
        if selected_value == 'bird':
            return {'display': 'block'}, {'display': 'none'}
        elif selected_value == 'vitamins':
            return {'display': 'none'}, {'display': 'block'}
        

    @dash_app.callback(
        [Output('Feed_req', 'children'),
        Output('feedNeed_store', 'data')],
        [Input('stages_dd', 'value'),
        Input('nBird', 'value')]
    )
        
    def feedRequired(stages, n_bird): 
        if n_bird is None:
            feedNeeded = 0
            reqDay = ''
        else:
            if stages == "Pre-starter":
                feedNeeded = 0.5 * n_bird
                reqDay = 'Day 0-7'
            elif stages == 'Starter':  
                feedNeeded = 1 * n_bird
                reqDay = 'Day 8-21'
            elif stages == 'Finisher':  
                feedNeeded = 2.8 * n_bird
                reqDay = 'Day 22-42'
            else:
                feedNeeded = 0 
                reqDay = '' 

        feedRequiredDisplay = [
            html.P(html.B("Feed Req.")),
            html.P(html.B(f"{round(feedNeeded, 2)} KG", style={"font-size": "13px", "position": "relative", 'bottom': '5px', "margin-right": "2px", "margin-left": "2px"})),
            html.P(html.I(reqDay, style={"font-size": "11px", "position": "relative", 'bottom': '15px', "margin-right": "2px", "margin-left": "2px"}))
        ]
        return feedRequiredDisplay, feedNeeded



    @dash_app.callback(
        [Output('ingredient_table', 'data'),
        Output('report_table', 'data'),
        Output('nutrient_table', 'data'),
        Output('ME_CP', 'children'),
        Output('Total_Cost', 'children'),
        Output('COST_25kg', 'children')],
        [Input('ingredient_dd', 'value'),
        Input('ingredient_table', 'data_timestamp'),
        Input('feedNeed_store', 'data')],
        [State('ingredient_table', 'data')]
    )
    def update_table(selected_ingredients, _, feedNeed, current_data):
        # Initialize DataFrame with current data
        if current_data:
            df_output = pd.DataFrame(current_data)
        else:
            df_output = pd.DataFrame(columns=columns)

        # Remove any existing 'Total' row
        df_output = df_output[df_output['INGREDIENT'] != 'Total']

        # Remove ingredients that are no longer selected
        if selected_ingredients is not None:
            df_output = df_output[df_output['INGREDIENT'].isin(selected_ingredients)]

        # Add rows for selected ingredients if not already present
        if selected_ingredients:
            for ingredient in selected_ingredients:
                if not df_output[df_output['INGREDIENT'] == ingredient].empty:
                    continue
                new_row = pd.DataFrame([{col: '' for col in columns}])
                new_row.at[0, 'INGREDIENT'] = ingredient
                new_row.at[0, 'PRICE/KG'] = 0
                df_output = pd.concat([df_output, new_row], ignore_index=True)

        # Calculate values for each ingredient
        for i, row in df_output.iterrows():
            if row['INGREDIENT'] in dfx['INGREDIENT'].values:
                ing_values = dfx[dfx['INGREDIENT'] == row['INGREDIENT']].iloc[0]
                try:
                    quantity = float(row['QUANTITY']) if row['QUANTITY'] else 0
                    price_per_kg = float(row['PRICE/KG']) if row['PRICE/KG'] != '' else 0
                except ValueError:
                    quantity = 0
                    price_per_kg = 0

                if quantity > 0:
                    for col in columns[2:]:
                        if col in df_output.columns and col in ing_values.index:
                            df_output.at[i, col] = round((ing_values[col] / 100) * quantity, 4)
                    df_output.at[i, 'QUANTITY PRICE'] = quantity * price_per_kg
                else:
                    for col in columns[2:]:
                        df_output.at[i, col] = ''
                    df_output.at[i, 'QUANTITY PRICE'] = ''

        # Ensure numeric columns are correctly converted to numeric types
        numeric_cols = ['QUANTITY', 'PRICE/KG', 'QUANTITY PRICE', 'DRY MATTER', 'CP', 'FAT', 'FIBRE', 'CAL.', 
                        'PHOS.TOTAL', 'AVAIL PHOS', 'ME/POULT', 'ME/SWINE', 'METH', 'CYSTINE', 'METH+CYST',
                        'LYSINE', 'TRYPTOPHAN', 'THREONINE', 'VIT A IU/GM', 'VIT D3 IU/GM',
                        'VIT E IU/GM', 'RIBOFLAVIN', 'PANTO ACID', 'CHOLINE', 'B 12', 'NIACIN',
                        'XANTHOPYL', 'SALT', 'SODIUM', 'POTASSIUM', 'MAGNESIUM', 'SULPHUR',
                        'MANGANESE', 'IRON', 'COPPER', 'ZINC', 'SELENIUM', 'IODINE', 'LINOLEIC A']

        df_output[numeric_cols] = df_output[numeric_cols].apply(pd.to_numeric, errors='coerce').fillna(0)

        # Calculate totals for each numeric column
        total_row = round(df_output[numeric_cols].sum().to_frame().T, 4)

        # Set the total row values
        total_quantity = df_output['QUANTITY'].sum()
        total_quantity_price = df_output['QUANTITY PRICE'].sum()

        total_row['INGREDIENT'] = 'Total'
        total_row['QUANTITY'] = total_quantity
        total_row['QUANTITY PRICE'] = total_quantity_price

        # Calculate the weighted average PRICE/KG
        if total_quantity > 0:
            total_row['PRICE/KG'] = total_quantity_price / total_quantity
        else:
            total_row['PRICE/KG'] = 0

        # Append the total row only once
        df_output = pd.concat([df_output, total_row], ignore_index=True)

        # Calculate total ME / total CP
        total_me = total_row['ME/POULT'].values[0]
        total_cp = total_row['CP'].values[0]
        me_cp_value = total_me / total_cp if total_cp > 0 else 0

        # Calculate COST/25kg
        cost_25kg = total_row['PRICE/KG'].values[0] * 25

        # Prepare the report table data
        report_data = df_output[['INGREDIENT', 'PRICE/KG', 'QUANTITY', 'QUANTITY PRICE']].to_dict('records')

        # Adjust the report table quantity using feedNeed
        for item in report_data:
            if item['INGREDIENT'] != 'Total' and item['QUANTITY']:
                item['QUANTITY'] = (item['QUANTITY'] / 100) * feedNeed
            item['AMOUNT'] = item['QUANTITY'] * item['PRICE/KG'] if item['QUANTITY'] and item['PRICE/KG'] else 0

        # Calculate the total amount based on the adjusted quantities
        total_amount = sum(item['AMOUNT'] for item in report_data if item['INGREDIENT'] != 'Total')
        total_quantity_report = sum(item['QUANTITY'] for item in report_data if item['INGREDIENT'] != 'Total')

        # Set the correct total row values for the report table
        for row in report_data:
            if row['INGREDIENT'] == 'Total':
                row['QUANTITY'] = total_quantity_report
                row['QUANTITY PRICE'] = total_quantity_price
                row['AMOUNT'] = total_amount  # Sum of all individual amounts
                break
        else:
            report_data.append({
                'INGREDIENT': 'Total',
                'PRICE/KG': '',
                'QUANTITY': total_quantity_report,
                'QUANTITY PRICE': total_quantity_price,
                'AMOUNT': total_amount  # Add total amount
            })

        # Prepare the nutrient composition data, excluding specified columns
        exclude_cols = ['PRICE/KG', 'QUANTITY', 'QUANTITY PRICE', 'DRY MATTER']
        nutrient_data = []
        nutrient_totals = total_row
        for col in numeric_cols:
            if col not in exclude_cols and nutrient_totals[col].values[0] > 0:
                nutrient_data.append({'Nutrient': col, 'Actual': nutrient_totals[col].values[0]})

        # Update ME/CP, Total Cost, and COST/25kg values
        return (df_output.to_dict('records'),
                report_data,
                nutrient_data,
                [html.P(html.B("ME/CP(KG)")), html.P(html.B(f"{me_cp_value:.2f}"))],  # Update ME/CP
                [html.P(html.B("TOTAL COST")), html.P(html.B(f"{total_amount:.2f}"))],  # Update Total Cost
                [html.P(html.B("COST/25kg")), html.P(html.B(f"{cost_25kg:.2f}"))])  # Update COST/25kg

    import os
    import sqlite3
    from datetime import datetime
    import pandas as pd
    from dash import dcc, html
    from dash.dependencies import Input, Output, State
    from flask_login import login_required
    from docx import Document
    
    # Initialize SQLite3 database
    DB_PATH = os.path.join(os.path.expanduser("~"), "Documents", "feed_analysis.db")
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    
    # Create necessary tables if they don't exist
    def init_db():
        with sqlite3.connect(DB_PATH) as conn:
            cursor = conn.cursor()
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS feeds (
                    feed_code TEXT PRIMARY KEY,
                    feed_name TEXT,
                    report_date TEXT
                )
            """)
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS ingredients (
                    ingredient_name TEXT PRIMARY KEY
                )
            """)
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS feed_ingredients (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    feed_code TEXT,
                    ingredient_name TEXT,
                    price_per_kg REAL,
                    quantity REAL,
                    quantity_price REAL,
                    amount REAL,
                    FOREIGN KEY (feed_code) REFERENCES feeds(feed_code),
                    FOREIGN KEY (ingredient_name) REFERENCES ingredients(ingredient_name)
                )
            """)
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS nutrient_composition (
                    feed_code TEXT PRIMARY KEY,
                    cp REAL, fat REAL, fibre REAL, cal REAL,
                    phos_total REAL, avail_phos REAL, me_poult REAL, me_swine REAL,
                    meth REAL, cystine REAL, meth_cyst REAL, lysine REAL,
                    tryptophan REAL, threonine REAL, vit_a_iu_gm REAL, vit_e_iu_gm REAL,
                    riboflavin REAL, panto_acid REAL, choline REAL, niacin REAL,
                    sodium REAL, potassium REAL, magnesium REAL, sulphur REAL,
                    manganese REAL, iron REAL, copper REAL, zinc REAL,
                    selenium REAL, iodine REAL
                )
            """)
    
        init_db()

    # Callback to generate and download the Word document
    @dash_app.callback(
        Output("download_word", "data"),
        [Input("download_word_button", "n_clicks")],
        [State("Company_name", "value"),
         State("feed_name", "value"),
         State("feed_code", "children"),
         State("report_table", "data"),
         State("nutrient_table", "data")]
    )
    def generate_word_document(n_clicks, company_name, feed_name, feed_code, report_data, nutrient_data):
        if n_clicks > 0:
            # **Create Word Document**
            doc = Document()
            doc.add_heading(f"{company_name} - Feed Analysis Report", level=1)
            doc.add_paragraph(f"Feed Name: {feed_name}")
            doc.add_paragraph(f"Feed Code: {feed_code}")
            doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
            # **Add Report Data**
            if report_data:
                doc.add_heading("Report Table", level=2)
                table = doc.add_table(rows=1, cols=len(report_data[0].keys()))
                hdr_cells = table.rows[0].cells
                for i, key in enumerate(report_data[0].keys()):
                    hdr_cells[i].text = key
    
                for row in report_data:
                    row_cells = table.add_row().cells
                    for i, key in enumerate(row.keys()):
                        row_cells[i].text = str(row[key])
    
            # **Add Nutrient Data**
            if nutrient_data:
                doc.add_heading("Nutrient Composition", level=2)
                table = doc.add_table(rows=1, cols=len(nutrient_data[0].keys()))
                hdr_cells = table.rows[0].cells
                for i, key in enumerate(nutrient_data[0].keys()):
                    hdr_cells[i].text = key
    
                for row in nutrient_data:
                    row_cells = table.add_row().cells
                    for i, key in enumerate(row.keys()):
                        row_cells[i].text = str(row[key])
    
            # **Save the Document**
            file_name = "Feed Analysis Report.docx"
            doc.save(file_name)
    
            # **Insert Data into SQLite Database**
            with sqlite3.connect(DB_PATH) as conn:
                cursor = conn.cursor()
    
                # Insert Feed Details
                cursor.execute("""
                    INSERT OR REPLACE INTO feeds (feed_code, feed_name, report_date)
                    VALUES (?, ?, ?)
                """, (feed_code, feed_name, datetime.now().strftime('%Y-%m-%d %H:%M:%S')))
    
                # Insert Ingredients Data
                if report_data:
                    for row in report_data:
                        ingredient_name = row['INGREDIENT']
                        cursor.execute("""
                            INSERT OR IGNORE INTO ingredients (ingredient_name)
                            VALUES (?)
                        """, (ingredient_name,))
                        cursor.execute("""
                            INSERT INTO feed_ingredients (feed_code, ingredient_name, price_per_kg, quantity, quantity_price, amount)
                            VALUES (?, ?, ?, ?, ?, ?)
                        """, (
                            feed_code, ingredient_name, row['PRICE/KG'], row['QUANTITY'], 
                            row['QUANTITY PRICE'], row['AMOUNT']
                        ))
    
                # Insert Nutrient Composition
                if nutrient_data:
                    nutrient_dict = {item['Nutrient']: item['Actual'] for item in nutrient_data}
                    cursor.execute("""
                        INSERT OR REPLACE INTO nutrient_composition (
                            feed_code, cp, fat, fibre, cal, phos_total, avail_phos, me_poult, me_swine, meth, cystine,
                            meth_cyst, lysine, tryptophan, threonine, vit_a_iu_gm, vit_e_iu_gm, riboflavin, panto_acid,
                            choline, niacin, sodium, potassium, magnesium, sulphur, manganese, iron, copper, zinc,
                            selenium, iodine
                        )
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (
                        feed_code, nutrient_dict.get("CP", 0), nutrient_dict.get("FAT", 0), nutrient_dict.get("FIBRE", 0),
                        nutrient_dict.get("CAL", 0), nutrient_dict.get("PHOS.TOTAL", 0), nutrient_dict.get("AVAIL PHOS", 0),
                        nutrient_dict.get("ME/POULT", 0), nutrient_dict.get("ME/SWINE", 0), nutrient_dict.get("METH", 0),
                        nutrient_dict.get("CYSTINE", 0), nutrient_dict.get("METH+CYST", 0), nutrient_dict.get("LYSINE", 0),
                        nutrient_dict.get("TRYPTOPHAN", 0), nutrient_dict.get("THREONINE", 0), nutrient_dict.get("VIT A IU/GM", 0),
                        nutrient_dict.get("VIT E IU/GM", 0), nutrient_dict.get("RIBOFLAVIN", 0), nutrient_dict.get("PANTO ACID", 0),
                        nutrient_dict.get("CHOLINE", 0), nutrient_dict.get("NIACIN", 0), nutrient_dict.get("SODIUM", 0),
                        nutrient_dict.get("POTASSIUM", 0), nutrient_dict.get("MAGNESIUM", 0), nutrient_dict.get("SULPHUR", 0),
                        nutrient_dict.get("MANGANESE", 0), nutrient_dict.get("IRON", 0), nutrient_dict.get("COPPER", 0),
                        nutrient_dict.get("ZINC", 0), nutrient_dict.get("SELENIUM", 0), nutrient_dict.get("IODINE", 0)
                    ))
    
            # **Send the Word Document for Download**
            return dcc.send_file(file_name)
    
        # **Return No Update if No Clicks**
        return dash.no_update
    
    
    # **Secure Dash Views with Login**
    for view_function in dash_app.server.view_functions:
        if view_function.startswith(dash_app.config.url_base_pathname):
            dash_app.server.view_functions[view_function] = login_required(
                dash_app.server.view_functions[view_function]
            )
    
    # **Return the Dash App**
    return dash_app
